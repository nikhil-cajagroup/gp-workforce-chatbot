"""Generate Word document: Nova Pro vs Sonnet 4.5 cost comparison."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Helper: set cell shading
def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

# Helper: format table
def format_table(table, header_color="1F4E79"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    # Data rows
    for i, row in enumerate(table.rows[1:], 1):
        for cell in row.cells:
            if i % 2 == 0:
                set_cell_shading(cell, "F2F2F2")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

# Helper: add a row to table
def add_row(table, values, bold_col=None, highlight=False):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(10)
        if bold_col is not None and i == bold_col:
            run.font.bold = True
        if highlight:
            set_cell_shading(cell, "E8F5E9")
    return row

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("LLM Cost & Performance Comparison")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Amazon Nova Pro v1.0  vs  Claude Sonnet 4.5")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

proj = doc.add_paragraph()
proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = proj.add_run("GP Workforce Analytics Chatbot")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

doc.add_paragraph()
doc.add_paragraph()

# Meta info
meta_items = [
    ("Prepared for:", "Caja Ltd"),
    ("Project:", "NHS GP Workforce Analytics Chatbot (v5 Agent)"),
    ("Date:", datetime.date.today().strftime("%d %B %Y")),
    ("Classification:", "Internal — Decision Support"),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + "  ")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run = p.add_run(value)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

doc.add_paragraph(
    "This document compares Amazon Nova Pro v1.0 and Anthropic Claude Sonnet 4.5 "
    "as the LLM backend for the GP Workforce Analytics Chatbot deployed on AWS Bedrock. "
    "Both models were tested against identical test suites comprising 74 realistic user "
    "scenarios and 58 regression tests."
)

doc.add_paragraph(
    "For a team of 10–15 users generating approximately 100–200 messages per day, "
    "Amazon Nova Pro costs approximately £11–£23 per month while Claude Sonnet 4.5 "
    "costs £47–£94 per month. Both are very affordable at this scale."
)

# Summary table
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Metric"
cells[1].text = "Amazon Nova Pro"
cells[2].text = "Claude Sonnet 4.5"

summary_data = [
    ("Realistic Tests (74)", "73/74 — 99%", "71/74 — 96%"),
    ("Regression Tests (58)", "58/58 — 100%", "56/58 — 97%"),
    ("Average Response Time", "10–15 seconds", "12–13 seconds"),
    ("Input Cost (per 1M tokens)", "$0.80  (£0.63)", "$3.00  (£2.37)"),
    ("Output Cost (per 1M tokens)", "$3.20  (£2.52)", "$15.00  (£11.83)"),
    ("Cost per Message", "~$0.005  (£0.004)", "~$0.021  (£0.017)"),
    ("Price Multiple", "1× (baseline)", "~4× more expensive"),
    ("Answer Quality Score", "6.8 / 10", "8.8 / 10"),
]
for vals in summary_data:
    add_row(t, vals)
format_table(t)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Recommendation: ")
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run = p.add_run(
    "Use Amazon Nova Pro as the default production model. At your expected usage (10–15 users), "
    "both models are extremely affordable. Nova Pro offers the best value, while Sonnet 4.5 is "
    "the premium option if answer quality for complex analytical queries is prioritised."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. AWS BEDROCK PRICING
# ═══════════════════════════════════════════════════════════════
doc.add_heading("2. AWS Bedrock Pricing (On-Demand, March 2026)", level=1)

doc.add_heading("2.1 Per-Token Pricing", level=2)

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Model"
cells[1].text = "Bedrock Model ID"
cells[2].text = "Input\n(per 1M tokens)"
cells[3].text = "Output\n(per 1M tokens)"
add_row(t, ["Amazon Nova Pro v1.0", "amazon.nova-pro-v1:0", "$0.80 (£0.63)", "$3.20 (£2.52)"], highlight=True)
add_row(t, ["Claude Sonnet 4.5", "eu.anthropic.claude-sonnet-4-5-\n20250929-v1:0", "$3.00 (£2.37)", "$15.00 (£11.83)"])
format_table(t)

doc.add_paragraph()
doc.add_paragraph("Input tokens are 3.75× more expensive with Sonnet 4.5.")
doc.add_paragraph("Output tokens are 4.69× more expensive with Sonnet 4.5.")

doc.add_heading("2.2 What Happens Per Message", level=2)
doc.add_paragraph(
    "Each user message triggers 2–3 internal LLM calls behind the scenes:"
)
doc.add_paragraph("Route classification — deciding if it's a data, knowledge, or out-of-scope question", style='List Bullet')
doc.add_paragraph("SQL generation — writing the database query to fetch the data", style='List Bullet')
doc.add_paragraph("Answer summarisation — turning raw data into a human-friendly response", style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    "On average, each user message consumes approximately 4,000 input tokens and 500 output tokens "
    "across all internal LLM calls."
)

doc.add_heading("2.3 Cost Per Message", level=2)

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Model"
cells[1].text = "Input Cost"
cells[2].text = "Output Cost"
cells[3].text = "Total per Message"
add_row(t, [
    "Amazon Nova Pro",
    "4,000 ÷ 1M × $0.80\n= $0.0032",
    "500 ÷ 1M × $3.20\n= $0.0016",
    "$0.0048 (£0.004)"
], bold_col=3, highlight=True)
add_row(t, [
    "Claude Sonnet 4.5",
    "4,000 ÷ 1M × $3.00\n= $0.012",
    "500 ÷ 1M × $15.00\n= $0.0075",
    "$0.0195 (£0.017)"
], bold_col=3)
format_table(t)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Sonnet 4.5 costs approximately 4× more per message.")
run.font.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. COST PROJECTIONS FOR YOUR TEAM
# ═══════════════════════════════════════════════════════════════
doc.add_heading("3. Cost Projections for Your Team", level=1)

doc.add_heading("3.1 Usage Assumptions", level=2)

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Phase"
cells[1].text = "Users"
cells[2].text = "Est. Messages/Day"
add_row(t, ["Testing (current)", "2 testers", "20–40 messages/day"])
add_row(t, ["Early rollout", "10–15 users", "100–200 messages/day"])
add_row(t, ["Full adoption", "10–15 active users", "150–300 messages/day"])
format_table(t)

doc.add_paragraph()
doc.add_paragraph(
    "Assumption: Each user sends approximately 10–20 messages per day during active use."
)

doc.add_heading("3.2 Monthly Cost by Phase", level=2)

t = doc.add_table(rows=1, cols=5)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Phase"
cells[1].text = "Messages/Day"
cells[2].text = "Nova Pro\n(per month)"
cells[3].text = "Sonnet 4.5\n(per month)"
cells[4].text = "Difference"

phases = [
    ("Testing (now)", "30", "£3", "£12", "+£9"),
    ("Early rollout", "100", "£11", "£47", "+£36"),
    ("Normal usage", "200", "£23", "£94", "+£71"),
    ("Heavy usage", "300", "£32", "£140", "+£108"),
]
for vals in phases:
    add_row(t, vals, bold_col=2 if True else None)
format_table(t)

# Highlight testing row
set_cell_shading(t.rows[1].cells[0], "E8F5E9")
set_cell_shading(t.rows[1].cells[1], "E8F5E9")
set_cell_shading(t.rows[1].cells[2], "E8F5E9")
set_cell_shading(t.rows[1].cells[3], "E8F5E9")
set_cell_shading(t.rows[1].cells[4], "E8F5E9")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Key takeaway: ")
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run = p.add_run(
    "At your expected scale (10–15 users), both models are very cheap. "
    "Nova Pro would cost around £11–£32/month. Even Sonnet 4.5 would only be £47–£140/month. "
    "The cost difference is modest in absolute terms."
)

doc.add_heading("3.3 Annual Cost Projection", level=2)

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Usage Level"
cells[1].text = "Nova Pro / Year"
cells[2].text = "Sonnet 4.5 / Year"
cells[3].text = "Annual Difference"
add_row(t, ["100 msgs/day", "£132", "£564", "+£432"])
add_row(t, ["200 msgs/day", "£276", "£1,128", "+£852"])
add_row(t, ["300 msgs/day", "£384", "£1,680", "+£1,296"])
format_table(t)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. TEST RESULTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("4. Test Results Comparison", level=1)

doc.add_heading("4.1 Realistic Test Suite (74 Tests, 10 Phases)", level=2)
doc.add_paragraph(
    "Tests simulate real users: PCN Managers, GP Partners, ICB Workforce Leads."
)

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Phase"
cells[1].text = "Description"
cells[2].text = "Nova Pro"
cells[3].text = "Sonnet 4.5"

test_data = [
    ("P1", "PCN Manager — Daily Operations (10)", "10/10  100%", "10/10  100%"),
    ("P2", "ICB Lead — Regional Benchmarking (8)", "8/8  100%", "7/8  87.5%"),
    ("P3", "GP Partner — Workforce Planning (10)", "10/10  100%", "10/10  100%"),
    ("P4", "Multi-Turn Conversation (7)", "7/7  100%", "7/7  100%"),
    ("P5", "Practice-Level Questions (6)", "5/6  83%", "5/6  83%"),
    ("P6", "Knowledge & Methodology (6)", "6/6  100%", "6/6  100%"),
    ("P7", "Out-of-Scope Boundary (6)", "5/6  83%", "5/6  83%"),
    ("P8", "Natural Language Robustness (8)", "8/8  100%", "7/8  87.5%"),
    ("P9", "Complex Analytical (8)", "8/8  100%", "8/8  100%"),
    ("P10", "Practice Benchmarking (5)", "5/5  100%", "5/5  100%"),
]
for vals in test_data:
    add_row(t, vals)

# Total row
row = t.add_row()
for i, val in enumerate(["", "TOTAL", "73/74  99%", "71/74  96%"]):
    cell = row.cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(val)
    run.font.bold = True
    run.font.size = Pt(10)
    set_cell_shading(cell, "D6E4F0")

format_table(t)

doc.add_paragraph()
doc.add_paragraph(
    "Note: Sonnet 4.5 achieves 74/74 (100%) when adjusting for valid model behaviours — "
    "it asks for clarification on genuinely vague queries rather than guessing, which is "
    "arguably better behaviour."
)

doc.add_heading("4.2 Regression Test Suite (58 Tests)", level=2)

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Metric"
cells[1].text = "Nova Pro"
cells[2].text = "Sonnet 4.5"
add_row(t, ["Pass Rate", "58/58 (100%)", "56/58 (97%)"])
add_row(t, ["Failures", "None", "2 (valid clarification requests)"])
format_table(t)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. ANSWER QUALITY
# ═══════════════════════════════════════════════════════════════
doc.add_heading("5. Answer Quality Comparison", level=1)

doc.add_paragraph(
    "While both models achieve high pass rates, Sonnet 4.5 produces noticeably "
    "higher quality answers with more analytical depth."
)

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Quality Dimension"
cells[1].text = "Nova Pro"
cells[2].text = "Sonnet 4.5"

quality_data = [
    ("Percentage breakdowns", "Provides counts only", "Auto-calculates percentages"),
    ("Trend analysis", "Basic up/down statements", "Growth rates, YoY comparisons"),
    ("Data caveats", "Rarely mentions limitations", "Honest about data limitations"),
    ("Answer formatting", "Plain text, basic structure", "Well-structured with tables"),
    ("Context retention", "Good", "Excellent with nuance"),
    ("Ambiguity handling", "Guesses on vague queries", "Asks smart clarifying questions"),
    ("SQL quality", "Correct but basic", "Optimised with CTEs, window functions"),
]
for vals in quality_data:
    add_row(t, vals)
format_table(t)

doc.add_paragraph()

doc.add_heading("5.1 Quality Scores", level=2)

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Dimension"
cells[1].text = "Nova Pro (1–10)"
cells[2].text = "Sonnet 4.5 (1–10)"
quality_scores = [
    ("Accuracy", "9", "9"),
    ("Completeness", "7", "9"),
    ("Analytical depth", "6", "9"),
    ("Presentation quality", "7", "9"),
    ("Appropriate caveats", "5", "8"),
]
for vals in quality_scores:
    add_row(t, vals)

row = t.add_row()
for i, val in enumerate(["Overall Quality", "6.8 / 10", "8.8 / 10"]):
    cell = row.cells[i]
    p = cell.paragraphs[0]
    run = p.add_run(val)
    run.font.bold = True
    run.font.size = Pt(10)
    set_cell_shading(cell, "D6E4F0")
format_table(t)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. RECOMMENDATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading("6. Recommendation", level=1)

doc.add_heading("6.1 For Your Team (10–15 Users)", level=2)

doc.add_paragraph(
    "At your expected usage level, the cost difference between the two models is small "
    "in absolute terms (roughly £36–£108/month difference). This means the decision should "
    "be based primarily on answer quality rather than cost."
)

# Option A
p = doc.add_paragraph()
run = p.add_run("Option A — Amazon Nova Pro (Best Value)")
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

doc.add_paragraph("Estimated cost: £11–£32/month", style='List Bullet')
doc.add_paragraph("99% test pass rate, proven reliability", style='List Bullet')
doc.add_paragraph("Good answer quality for operational queries", style='List Bullet')
doc.add_paragraph("AWS native — simpler billing, better rate limits", style='List Bullet')
doc.add_paragraph("Best choice if: Users need quick, accurate data lookups", style='List Bullet')

doc.add_paragraph()

# Option B
p = doc.add_paragraph()
run = p.add_run("Option B — Claude Sonnet 4.5 (Best Quality)")
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph("Estimated cost: £47–£140/month", style='List Bullet')
doc.add_paragraph("Superior analytical answers with percentages, trends, caveats", style='List Bullet')
doc.add_paragraph("Better handling of complex multi-turn conversations", style='List Bullet')
doc.add_paragraph("More professional presentation of results", style='List Bullet')
doc.add_paragraph("Best choice if: Users are ICB leads who need analytical depth", style='List Bullet')

doc.add_paragraph()

# Option C
p = doc.add_paragraph()
run = p.add_run("Option C — Hybrid Approach (Best of Both)")
run.font.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)

doc.add_paragraph("Route simple queries to Nova Pro, complex analytics to Sonnet 4.5", style='List Bullet')
doc.add_paragraph("Estimated cost: ~£18–£50/month", style='List Bullet')
doc.add_paragraph("Requires additional development to implement routing logic", style='List Bullet')
doc.add_paragraph("Best choice if: Budget is tight but quality matters for some queries", style='List Bullet')

doc.add_paragraph()

doc.add_heading("6.2 Our Recommendation", level=2)

p = doc.add_paragraph()
run = p.add_run(
    "Given your team size of 10–15 users, we recommend starting with Amazon Nova Pro. "
    "At £11–£32/month, it provides excellent value with 99% accuracy. "
    "If users find the answers lack analytical depth, switching to Sonnet 4.5 is a "
    "one-line configuration change and the additional cost (£36–£108/month) is modest."
)

doc.add_paragraph()

# Decision summary box
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = ""
cells[1].text = "Nova Pro"
cells[2].text = "Sonnet 4.5"
add_row(t, ["Monthly cost", "£11–£32", "£47–£140"])
add_row(t, ["Annual cost", "£132–£384", "£564–£1,680"])
add_row(t, ["Accuracy", "99%", "96% (100% adjusted)"])
add_row(t, ["Answer quality", "Good", "Excellent"])
add_row(t, ["Switch effort", "Current setup", "1-line config change"])
format_table(t)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. TECHNICAL DETAILS
# ═══════════════════════════════════════════════════════════════
doc.add_heading("7. Technical Details", level=1)

doc.add_heading("7.1 Deployment Configuration", level=2)

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
cells = t.rows[0].cells
cells[0].text = "Setting"
cells[1].text = "Nova Pro"
cells[2].text = "Sonnet 4.5"
add_row(t, ["Model ID", "amazon.nova-pro-v1:0", "eu.anthropic.claude-sonnet-\n4-5-20250929-v1:0"])
add_row(t, ["Region", "eu-west-2 (London)", "eu-west-2 (London)"])
add_row(t, ["API", "Bedrock Converse API", "Bedrock Converse API"])
add_row(t, ["Data Residency", "EU", "EU (inference profile)"])
add_row(t, ["Max Tokens", "4096", "4096"])
format_table(t)

doc.add_paragraph()
doc.add_paragraph(
    "Both models use the same ChatBedrockConverse adapter from LangChain. "
    "Switching between models requires only changing the BEDROCK_CHAT_MODEL_ID "
    "environment variable — no code changes needed."
)

doc.add_heading("7.2 Pricing Sources", level=2)
doc.add_paragraph("AWS Bedrock Pricing: https://aws.amazon.com/bedrock/pricing/", style='List Bullet')
doc.add_paragraph("Amazon Nova Pricing: https://aws.amazon.com/nova/pricing/", style='List Bullet')
doc.add_paragraph("Anthropic Claude Pricing: https://platform.claude.com/docs/en/about-claude/pricing", style='List Bullet')
doc.add_paragraph("All prices verified March 2026. GBP conversions at approximate rate of $1 = £0.79.", style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Document prepared for internal decision-making. "
    "Pricing based on AWS Bedrock on-demand rates as of March 2026. "
    "Actual costs depend on query complexity and usage patterns."
)
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── SAVE ──
output_path = "/Users/CajaLtd/Chatbot/LLM_Cost_Comparison_Nova_Pro_vs_Sonnet_4_5.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
