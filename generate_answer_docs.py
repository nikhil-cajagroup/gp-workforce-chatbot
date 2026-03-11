"""
Generate Word documents showing chatbot answers for team review.
Creates one doc per model with all questions and full answers.
"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_border_bottom(paragraph, color="1F4E79"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_doc(json_file, output_file, model_name, accent_color):
    with open(json_file) as f:
        data = json.load(f)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ═══ TITLE PAGE ═══
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GP Workforce Chatbot")
    run.font.size = Pt(28)
    run.font.color.rgb = accent_color
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Answer Quality Review")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    model_p = doc.add_paragraph()
    model_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = model_p.add_run(f"Model: {model_name}")
    run.font.size = Pt(18)
    run.font.color.rgb = accent_color
    run.font.bold = True

    doc.add_paragraph()

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(f"Test Date: {data['timestamp']}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    meta_p2 = doc.add_paragraph()
    meta_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p2.add_run(f"{len(data['results'])} questions tested")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    doc.add_paragraph()

    # Instructions
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("For Team Review")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Please review each answer for accuracy, clarity, and usefulness.\n"
        "Rate each answer from 1 (poor) to 5 (excellent) in the space provided.\n"
        "Compare with the other model's document to help decide which to use."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ═══ SUMMARY TABLE ═══
    doc.add_heading("Quick Summary", level=1)

    total_time = sum(r["elapsed"] for r in data["results"])
    avg_time = total_time / len(data["results"])
    answered = sum(1 for r in data["results"] if r.get("sql") or len(r.get("answer", "")) > 50)

    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    summary_items = [
        ("Model", model_name),
        ("Questions Tested", str(len(data["results"]))),
        ("Average Response Time", f"{avg_time:.1f} seconds"),
        ("Total Test Duration", f"{total_time:.0f}s ({total_time/60:.1f} min)"),
    ]
    for i, (label, value) in enumerate(summary_items):
        t.rows[i].cells[0].paragraphs[0].add_run(label).font.bold = True
        t.rows[i].cells[1].paragraphs[0].add_run(value)
        set_cell_shading(t.rows[i].cells[0], "E8EEF4")

    doc.add_paragraph()
    doc.add_page_break()

    # ═══ QUESTIONS AND ANSWERS ═══
    current_category = ""
    q_num = 0

    for result in data["results"]:
        q_num += 1
        category = result["category"]

        # New category header
        if category != current_category:
            current_category = category
            if q_num > 1:
                doc.add_paragraph()  # spacing
            h = doc.add_heading(category, level=1)
            for run in h.runs:
                run.font.color.rgb = accent_color

        # Question box
        q_para = doc.add_paragraph()
        run = q_para.add_run(f"Q{q_num}: ")
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = accent_color
        run = q_para.add_run(result["question"])
        run.font.size = Pt(12)
        run.font.bold = True

        # Response time
        time_p = doc.add_paragraph()
        run = time_p.add_run(f"Response time: {result['elapsed']}s")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

        # Answer
        answer = result.get("answer", "No answer returned")
        if not answer.strip():
            answer = "No answer returned"

        ans_label = doc.add_paragraph()
        run = ans_label.add_run("Chatbot Answer:")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        # Split answer into paragraphs for readability
        answer_lines = answer.split('\n')
        for line in answer_lines:
            if line.strip():
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(10)
                # indent slightly
                p.paragraph_format.left_indent = Cm(0.5)

        # SQL (if any)
        sql = result.get("sql", "")
        if sql and len(sql) > 10:
            sql_label = doc.add_paragraph()
            run = sql_label.add_run("SQL Generated:")
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            sql_p = doc.add_paragraph()
            run = sql_p.add_run(sql[:500])
            run.font.size = Pt(8)
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            sql_p.paragraph_format.left_indent = Cm(0.5)

        # Rating box
        doc.add_paragraph()
        rating_p = doc.add_paragraph()
        run = rating_p.add_run("Your Rating (1-5):  ☐ 1   ☐ 2   ☐ 3   ☐ 4   ☐ 5")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        notes_p = doc.add_paragraph()
        run = notes_p.add_run("Notes: _______________________________________________")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

        # Separator line
        sep = doc.add_paragraph()
        add_border_bottom(sep, "CCCCCC")

    # ═══ FINAL PAGE ═══
    doc.add_page_break()
    doc.add_heading("Overall Assessment", level=1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Model Reviewed: {model_name}")
    run.font.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    overall_items = [
        "Overall accuracy of answers (1-5): ___",
        "Clarity and readability (1-5): ___",
        "Analytical depth / usefulness (1-5): ___",
        "Handling of follow-up questions (1-5): ___",
        "Handling of informal / slang queries (1-5): ___",
    ]
    for item in overall_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Would you recommend this model for production use?")
    run.font.bold = True
    doc.add_paragraph("☐ Yes, definitely")
    doc.add_paragraph("☐ Yes, with minor improvements")
    doc.add_paragraph("☐ Not sure — need more testing")
    doc.add_paragraph("☐ No — prefer the other model")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Additional comments:")
    run.font.bold = True
    for _ in range(4):
        doc.add_paragraph("_" * 80)

    # Save
    doc.save(output_file)
    print(f"  ✅ Saved: {output_file}")


# ═══ GENERATE BOTH DOCS ═══
print("=" * 60)
print("  Generating Word Documents for Team Review")
print("=" * 60)

generate_doc(
    "/Users/CajaLtd/Chatbot/comparison_results_claude_sonnet_4.5.json",
    "/Users/CajaLtd/Chatbot/Chatbot_Answers_Claude_Sonnet_4_5.docx",
    "Claude Sonnet 4.5",
    RGBColor(0x1F, 0x4E, 0x79)  # Dark blue
)

generate_doc(
    "/Users/CajaLtd/Chatbot/comparison_results_amazon_nova_pro.json",
    "/Users/CajaLtd/Chatbot/Chatbot_Answers_Amazon_Nova_Pro.docx",
    "Amazon Nova Pro",
    RGBColor(0xE8, 0x6C, 0x00)  # Orange (AWS color)
)

print("\n" + "=" * 60)
print("  DONE — Both documents ready for team review")
print("=" * 60)
