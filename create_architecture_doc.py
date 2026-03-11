"""
Generate Architecture & Costing Word document for GP Workforce Chatbot production deployment.
"""
import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NHS_BLUE = RGBColor(0x00, 0x5E, 0xB8)
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GREY = RGBColor(0x99, 0x99, 0x99)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
ORANGE = RGBColor(0xE8, 0x6C, 0x00)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def add_styled_row(table, row_idx, cells_data, header=False, shade=None):
    row = table.rows[row_idx]
    for i, (text, bold, color) in enumerate(cells_data):
        p = row.cells[i].paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9 if not header else 9)
        run.font.bold = bold or header
        if color:
            run.font.color.rgb = color
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if header:
        for cell in row.cells:
            set_cell_shading(cell, "1F4E79")
    elif shade:
        for cell in row.cells:
            set_cell_shading(cell, shade)

def add_border_bottom(paragraph, color="1F4E79"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ═══════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("GP Workforce Chatbot")
run.font.size = Pt(28)
run.font.color.rgb = NHS_BLUE
run.font.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Production Architecture & Costing")
run.font.size = Pt(22)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("Deployment Plan for AWS with WordPress Integration")
run.font.size = Pt(14)
run.font.color.rgb = GREY

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Prepared: March 2026")
run.font.size = Pt(12)
run.font.color.rgb = LIGHT_GREY

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run("Current Users: 15 | Target: 15-30 users")
run.font.size = Pt(12)
run.font.color.rgb = LIGHT_GREY

doc.add_paragraph()
doc.add_paragraph()

scope = doc.add_paragraph()
scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = scope.add_run(
    "This document covers the recommended AWS architecture, production readiness\n"
    "review, step-by-step deployment plan, and monthly cost estimates for deploying\n"
    "the GP Workforce Chatbot behind the existing WordPress site."
)
run.font.size = Pt(11)
run.font.color.rgb = GREY

doc.add_page_break()

# ═══════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════
h = doc.add_heading("Contents", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

toc_items = [
    "1. Executive Summary",
    "2. Current State & Requirements",
    "3. AWS Services Comparison",
    "4. Recommended Architecture",
    "5. Production Readiness Review",
    "6. Deployment Plan (5 Phases)",
    "7. Monthly Cost Breakdown",
    "8. WordPress Integration",
    "9. Future Scaling Path",
    "10. Timeline & Next Steps",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════
h = doc.add_heading("1. Executive Summary", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
run = p.add_run(
    "The GP Workforce Chatbot is a FastAPI + React application that queries NHS workforce data "
    "via AWS Athena and generates natural language answers using AWS Bedrock LLMs. "
    "It currently runs locally and needs to be deployed to production behind the existing "
    "WordPress site on AWS."
)
run.font.size = Pt(11)

doc.add_paragraph()

# Key facts table
t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'
facts = [
    ("Current Stack", "FastAPI (Python) + React + AWS Athena + AWS Bedrock"),
    ("LLM Model", "Amazon Nova Pro (recommended) — £2/month for 15 users"),
    ("Current Users", "15 accounts with WordPress login"),
    ("Deployment Target", "AWS Elastic Beanstalk (Docker)"),
    ("WordPress Integration", "iframe embed behind existing login"),
    ("Estimated Monthly Cost", "£36 - £53/month (all AWS infrastructure)"),
    ("Deployment Timeline", "4-5 working days"),
]
for i, (label, value) in enumerate(facts):
    t.rows[i].cells[0].paragraphs[0].add_run(label).font.bold = True
    t.rows[i].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    t.rows[i].cells[1].paragraphs[0].add_run(value)
    t.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_shading(t.rows[i].cells[0], "E8EEF4")

doc.add_page_break()

# ═══════════════════════════════════════════
# 2. CURRENT STATE
# ═══════════════════════════════════════════
h = doc.add_heading("2. Current State & Requirements", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

doc.add_heading("What Exists Today", level=2)

existing = [
    "WordPress site with login system (15 user accounts)",
    "GP Practice Workforce dashboard with map visualisation and data tables",
    "Chatbot running locally (FastAPI backend + React frontend)",
    "AWS Athena queries against NHS GP workforce dataset",
    "AWS Bedrock for LLM inference (Nova Pro / Claude Sonnet)",
    "SSE streaming for real-time progress updates",
    "Multi-turn conversation with LangGraph checkpointing",
    "SQL injection protection, input validation, structured logging",
]
for item in existing:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_heading("What Needs to Happen", level=2)

needs = [
    "Deploy chatbot backend as a Docker container on AWS",
    "Host React frontend as static files (S3 + CloudFront)",
    "Add HTTPS, authentication, rate limiting",
    "Embed chatbot in WordPress page behind existing login",
    "Fix production readiness issues (see Section 5)",
    "Set up monitoring and cost alerts",
]
for item in needs:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 3. AWS SERVICES COMPARISON
# ═══════════════════════════════════════════
h = doc.add_heading("3. AWS Services Comparison", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
run = p.add_run("We evaluated four AWS compute options for hosting the FastAPI backend:")
run.font.size = Pt(11)

doc.add_paragraph()

# Comparison table
t = doc.add_table(rows=6, cols=5)
t.style = 'Table Grid'

headers = ["Criteria", "EC2 (raw)", "Elastic Beanstalk", "ECS Fargate", "Lambda"]
for i, h_text in enumerate(headers):
    p = t.rows[0].cells[i].paragraphs[0]
    run = p.add_run(h_text)
    run.font.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(t.rows[0].cells[i], "1F4E79")

rows_data = [
    ["Monthly Cost", "£7-15", "£7-15 (free wrapper)", "£15-30", "£3-10"],
    ["Setup Complexity", "High (manual Nginx, SSL, monitoring)", "Low (managed)", "Medium (task defs, ALB)", "High (cold starts, timeouts)"],
    ["Auto-restart on crash", "Manual (systemd)", "Built-in", "Built-in", "N/A"],
    ["SSL/HTTPS", "Manual (certbot)", "Built-in (ACM)", "Via ALB (ACM)", "Via API Gateway"],
    ["Best For", "Full control needed", "Small teams, quick deploy", "Microservices, auto-scaling", "Lightweight APIs"],
]

for r_idx, row_data in enumerate(rows_data, 1):
    for c_idx, text in enumerate(row_data):
        p = t.rows[r_idx].cells[c_idx].paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(8)
        if c_idx == 0:
            run.font.bold = True
    if r_idx % 2 == 0:
        for cell in t.rows[r_idx].cells:
            set_cell_shading(cell, "F5F5F5")

doc.add_paragraph()

# Recommendation box
rec = doc.add_paragraph()
run = rec.add_run("RECOMMENDATION: AWS Elastic Beanstalk")
run.font.bold = True
run.font.size = Pt(13)
run.font.color.rgb = GREEN

doc.add_paragraph()

reasons = [
    "Same cost as raw EC2 (Elastic Beanstalk itself is free — you only pay for EC2)",
    "Automatic Nginx reverse proxy, SSL via ACM, health monitoring",
    "Auto-restart on crash, rolling deployments, easy rollback",
    "Docker support out of the box — deploy with 'eb deploy'",
    "Industry standard for small-team Python/Docker deployments on AWS",
    "Easy upgrade path to larger instances or multiple instances if needed",
]
for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 4. RECOMMENDED ARCHITECTURE
# ═══════════════════════════════════════════
h = doc.add_heading("4. Recommended Architecture", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

doc.add_heading("Architecture Diagram", level=2)

# ASCII architecture diagram
arch = doc.add_paragraph()
run = arch.add_run(
    "WordPress Site (existing, login-protected)\n"
    "    |\n"
    "    +-- Chatbot page with <iframe>\n"
    "    |       |\n"
    "    |       v\n"
    "    |   CloudFront CDN --> S3 Bucket\n"
    "    |   (React frontend static files)\n"
    "    |       |\n"
    "    |       v (API calls via HTTPS)\n"
    "    |   Elastic Beanstalk\n"
    "    |   +-------------------------------+\n"
    "    |   | EC2 (t3.small)                |\n"
    "    |   | Nginx (auto-managed)          |\n"
    "    |   | Docker Container              |\n"
    "    |   |   FastAPI + Gunicorn          |\n"
    "    |   |   GP Workforce Agent v6.0     |\n"
    "    |   +-------------------------------+\n"
    "    |       |               |\n"
    "    |       v               v\n"
    "    |   AWS Athena      AWS Bedrock\n"
    "    |   (GP data)       (Nova Pro LLM)\n"
    "    |\n"
    "    +-- Rest of WordPress site (unchanged)\n"
)
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph()

doc.add_heading("How It Works", level=2)

flow = [
    "1. User logs into WordPress (existing login system, 15 accounts)",
    "2. User navigates to 'GP Workforce Chatbot' page",
    "3. WordPress page contains an <iframe> pointing to chatbot.yoursite.com",
    "4. React frontend loads from CloudFront/S3 (fast, cached globally)",
    "5. User types a question — React sends API request to Elastic Beanstalk",
    "6. FastAPI backend processes the question:",
    "   a. Generates SQL query using AWS Bedrock (Nova Pro LLM)",
    "   b. Executes query on AWS Athena (GP workforce dataset)",
    "   c. Formats results and generates natural language answer",
    "   d. Streams progress updates via SSE",
    "7. Answer displayed in chatbot with data table",
]
for item in flow:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(10)

doc.add_paragraph()

doc.add_heading("Key Design Decisions", level=2)

decisions = [
    ("WordPress login = access control", "No separate chatbot authentication needed. "
     "Only logged-in WordPress users can see the iframe page."),
    ("Elastic Beanstalk over raw EC2", "Free managed wrapper — same cost, less ops work. "
     "Handles Nginx, SSL, health checks, auto-restart."),
    ("CloudFront + S3 for frontend", "Static files served from edge locations. "
     "Fast globally, costs pennies per month."),
    ("Docker deployment", "Reproducible builds, easy rollback, works with EB natively."),
    ("Nova Pro over Claude Sonnet", "85% cheaper (£2 vs £13/month), 2.5x faster responses, "
     "comparable quality for structured data queries."),
]
for title_text, desc_text in decisions:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run(desc_text)
    run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════
# 5. PRODUCTION READINESS REVIEW
# ═══════════════════════════════════════════
h = doc.add_heading("5. Production Readiness Review", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

doc.add_heading("What's Already Production-Quality", level=2)

good_items = [
    "SQL injection protection (readonly enforcement, table whitelist, column validation)",
    "Structured logging with request IDs and tracing",
    "Request timeouts (90s default) with async wrapping",
    "Bounded caches (LRU) to prevent memory leaks",
    "Thread safety with locks on shared state",
    "Error sanitisation (no raw stack traces exposed to users)",
    "Input validation with Pydantic models (maxLength=1000)",
    "SSE streaming for real-time progress updates",
    "LangGraph checkpointing for conversation state persistence",
    "Semantic caching to reduce duplicate LLM calls",
]
for item in good_items:
    p = doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

doc.add_heading("Critical Issues (Must Fix Before Production)", level=2)

# Critical issues table
t = doc.add_table(rows=11, cols=3)
t.style = 'Table Grid'

crit_headers = ["#", "Issue", "Risk / Impact"]
for i, h_text in enumerate(crit_headers):
    p = t.rows[0].cells[i].paragraphs[0]
    run = p.add_run(h_text)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(t.rows[0].cells[i], "C62828")

critical_issues = [
    ("1", "No authentication on API", "Anyone can query Athena data and rack up LLM costs"),
    ("2", "No HTTPS", "Data transmitted in clear text"),
    ("3", "No rate limiting", "DDoS / cost explosion risk"),
    ("4", "Dev mode uvicorn (reload=True)", "Unstable in production, watches files"),
    ("5", "CORS allows only localhost", "WordPress domain will be blocked"),
    ("6", "AWS creds via shared file", "Should use IAM roles on EC2/EB"),
    ("7", "No Docker container", "Not reproducible, hard to deploy"),
    ("8", "Missing pip dependencies", "Install will fail on fresh server"),
    ("9", "SQLite checkpoint DB (514MB)", "Will fill disk, not scalable"),
    ("10", "Frontend API URL hardcoded", "Must point to production API"),
]
for r_idx, (num, issue, risk) in enumerate(critical_issues, 1):
    t.rows[r_idx].cells[0].paragraphs[0].add_run(num).font.size = Pt(9)
    t.rows[r_idx].cells[1].paragraphs[0].add_run(issue).font.size = Pt(9)
    t.rows[r_idx].cells[1].paragraphs[0].runs[0].font.bold = True
    t.rows[r_idx].cells[2].paragraphs[0].add_run(risk).font.size = Pt(9)
    if r_idx % 2 == 0:
        for cell in t.rows[r_idx].cells:
            set_cell_shading(cell, "FFF3F3")

doc.add_paragraph()

doc.add_heading("Important Issues (Should Fix)", level=2)

t = doc.add_table(rows=7, cols=2)
t.style = 'Table Grid'

imp_headers = ["Issue", "Recommendation"]
for i, h_text in enumerate(imp_headers):
    p = t.rows[0].cells[i].paragraphs[0]
    run = p.add_run(h_text)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(t.rows[0].cells[i], "E86C00")

important_issues = [
    ("No process manager", "Use gunicorn with uvicorn workers (handled by Docker CMD)"),
    ("No reverse proxy", "Elastic Beanstalk manages Nginx automatically"),
    ("No monitoring/alerting", "Add CloudWatch metrics + alarms"),
    ("Debug panel in production UI", "Hide or remove 'Show Debug' button with SQL"),
    ("Checkpoint DB cleanup", "Add TTL/rotation for sessions older than 7 days"),
    ("No .env.example file", "Document all required environment variables"),
]
for r_idx, (issue, rec_text) in enumerate(important_issues, 1):
    t.rows[r_idx].cells[0].paragraphs[0].add_run(issue).font.size = Pt(9)
    t.rows[r_idx].cells[0].paragraphs[0].runs[0].font.bold = True
    t.rows[r_idx].cells[1].paragraphs[0].add_run(rec_text).font.size = Pt(9)
    if r_idx % 2 == 0:
        for cell in t.rows[r_idx].cells:
            set_cell_shading(cell, "FFF8F0")

doc.add_page_break()

# ═══════════════════════════════════════════
# 6. DEPLOYMENT PLAN
# ═══════════════════════════════════════════
h = doc.add_heading("6. Deployment Plan (5 Phases)", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

# Phase 1
doc.add_heading("Phase 1: Fix Critical Code Issues (1-2 days)", level=2)

phase1 = [
    "Rename requirement.txt to requirements.txt and add missing dependencies "
    "(sse-starlette, numpy, langgraph, langgraph-checkpoint-sqlite, gunicorn)",
    "Remove reload=True from uvicorn.run() — add production startup via gunicorn",
    "Make CORS configurable via environment variable (accept WordPress domain)",
    "Add API key authentication middleware (X-API-Key header)",
    "Add rate limiting with slowapi (30 requests/min per session)",
    "Add deep health check (verify Athena + Bedrock connectivity)",
    "Add checkpoint DB cleanup (auto-prune sessions older than 7 days)",
    "Create .env.example with all required environment variables documented",
    "Hide/remove debug panel from production frontend build",
]
for item in phase1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Phase 2
doc.add_heading("Phase 2: Containerise with Docker (1 day)", level=2)

phase2 = [
    "Create Dockerfile (Python 3.11 slim, gunicorn + uvicorn workers)",
    "Create .dockerignore (exclude node_modules, __pycache__, test files, SQLite DB)",
    "Create docker-compose.yml for local testing",
    "Build React frontend for production (VITE_API_BASE=https://api.yoursite.com)",
    "Test Docker container locally — verify all endpoints work",
]
for item in phase2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Phase 3
doc.add_heading("Phase 3: AWS Infrastructure (1-2 days)", level=2)

phase3 = [
    "Create S3 bucket for React frontend static files",
    "Create CloudFront distribution pointing to S3 bucket",
    "Create Elastic Beanstalk Docker environment (t3.small, single instance)",
    "Configure EB environment: HTTPS via ACM certificate, environment variables",
    "Create IAM role for EB instance (Athena, S3, Bedrock access)",
    "Set up Route53 DNS: chatbot.yoursite.com (CloudFront), api.chatbot.yoursite.com (EB)",
    "Store API_KEY in EB environment variables or AWS Secrets Manager",
]
for item in phase3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# Phase 4
doc.add_heading("Phase 4: WordPress Integration (0.5 day)", level=2)

phase4 = [
    "Create new WordPress page: 'GP Workforce Chatbot'",
    "Add iframe embed pointing to https://chatbot.yoursite.com",
    "Test: verify CORS, SSL, streaming all work through iframe",
    "Verify only logged-in WordPress users can access the page",
]
for item in phase4:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run("\niframe code:")
run.font.bold = True
run.font.size = Pt(10)

code = doc.add_paragraph()
run = code.add_run(
    '<iframe\n'
    '  src="https://chatbot.yoursite.com"\n'
    '  width="100%" height="700px"\n'
    '  frameborder="0"\n'
    '  style="border-radius: 12px; box-shadow: 0 2px 10px rgba(0,0,0,0.1);"\n'
    '></iframe>'
)
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph()

# Phase 5
doc.add_heading("Phase 5: Monitoring & Go-Live (0.5 day)", level=2)

phase5 = [
    "Set up CloudWatch alarms (CPU > 80%, 5xx errors > 5%, response time > 30s)",
    "Enable CloudWatch Logs for EB container",
    "Set up AWS Billing cost alerts (Bedrock spend)",
    "Test with 2 team members — verify everything works end-to-end",
    "Roll out to all 15 users",
    "Monitor for 1 week, then remove debug panel from UI",
]
for item in phase5:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 7. MONTHLY COST BREAKDOWN
# ═══════════════════════════════════════════
h = doc.add_heading("7. Monthly Cost Breakdown", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
run = p.add_run("Estimated monthly costs for 15 users with Amazon Nova Pro:")
run.font.size = Pt(11)

doc.add_paragraph()

# Cost table
t = doc.add_table(rows=10, cols=3)
t.style = 'Table Grid'

cost_headers = ["AWS Service", "What It Does", "Monthly Cost"]
for i, h_text in enumerate(cost_headers):
    p = t.rows[0].cells[i].paragraphs[0]
    run = p.add_run(h_text)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(t.rows[0].cells[i], "1F4E79")

costs = [
    ("EC2 via Elastic Beanstalk", "t3.small (2 vCPU, 2GB RAM) — runs Docker container", "£12-15"),
    ("Elastic Beanstalk", "Manages Nginx, SSL, health checks, auto-restart", "Free"),
    ("CloudFront CDN", "Serves React frontend globally", "£1-3"),
    ("S3 (frontend)", "Stores React static files", "< £1"),
    ("Route53 DNS", "Custom domain routing", "£1"),
    ("ACM SSL Certificate", "HTTPS for both frontend and API", "Free"),
    ("AWS Bedrock (Nova Pro)", "LLM inference — 500 queries/day estimate", "£2-5"),
    ("AWS Athena", "SQL queries on GP workforce data", "£5-15"),
    ("CloudWatch", "Monitoring, logs, alarms", "£3-5"),
]

for r_idx, (service, desc, cost) in enumerate(costs, 1):
    t.rows[r_idx].cells[0].paragraphs[0].add_run(service).font.size = Pt(9)
    t.rows[r_idx].cells[0].paragraphs[0].runs[0].font.bold = True
    t.rows[r_idx].cells[1].paragraphs[0].add_run(desc).font.size = Pt(9)
    run = t.rows[r_idx].cells[2].paragraphs[0].add_run(cost)
    run.font.size = Pt(9)
    run.font.bold = True
    if r_idx % 2 == 0:
        for cell in t.rows[r_idx].cells:
            set_cell_shading(cell, "F5F5F5")

doc.add_paragraph()

# Total box
total = doc.add_paragraph()
run = total.add_run("TOTAL ESTIMATED: £36 - £53 per month")
run.font.bold = True
run.font.size = Pt(14)
run.font.color.rgb = GREEN

doc.add_paragraph()

# Annual
annual = doc.add_paragraph()
run = annual.add_run("Annual cost: £432 - £636 per year")
run.font.size = Pt(12)
run.font.color.rgb = DARK_BLUE

doc.add_paragraph()

# Comparison note
p = doc.add_paragraph()
run = p.add_run("Note: ")
run.font.bold = True
run.font.size = Pt(10)
run = p.add_run(
    "If using Claude Sonnet 4.5 instead of Nova Pro, the Bedrock cost increases from "
    "£2-5/month to £13-25/month, bringing the total to £47-68/month. "
    "Nova Pro is recommended for this use case (85% cheaper, 2.5x faster, comparable quality)."
)
run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════
# 8. WORDPRESS INTEGRATION
# ═══════════════════════════════════════════
h = doc.add_heading("8. WordPress Integration", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

doc.add_heading("Approach: iframe Embed (Recommended)", level=2)

p = doc.add_paragraph()
run = p.add_run(
    "The chatbot will be embedded as an iframe within a WordPress page. "
    "This is the simplest approach with zero impact on the existing WordPress site."
)
run.font.size = Pt(11)

doc.add_paragraph()

benefits = [
    "Zero changes to existing WordPress code or plugins",
    "Independent scaling — chatbot can be updated without touching WordPress",
    "Existing WordPress login controls who can see the page",
    "Clean separation of concerns (WordPress for content, chatbot on its own infrastructure)",
    "Easy to add/remove — just edit a WordPress page",
]
for item in benefits:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

doc.add_heading("How Users Will Experience It", level=2)

steps = [
    "1. User goes to your WordPress site and logs in (same as today)",
    "2. User clicks 'GP Workforce Chatbot' in the navigation menu",
    "3. The chatbot loads inside the page (via iframe) — looks like part of the site",
    "4. User asks questions, gets answers with data tables and charts",
    "5. Everything works within the existing WordPress page layout",
]
for step in steps:
    p = doc.add_paragraph()
    run = p.add_run(step)
    run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════
# 9. FUTURE SCALING
# ═══════════════════════════════════════════
h = doc.add_heading("9. Future Scaling Path", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
run = p.add_run(
    "The recommended architecture is designed to scale incrementally as user numbers grow:"
)
run.font.size = Pt(11)

doc.add_paragraph()

t = doc.add_table(rows=5, cols=3)
t.style = 'Table Grid'

scale_headers = ["Users", "Infrastructure", "Monthly Cost"]
for i, h_text in enumerate(scale_headers):
    p = t.rows[0].cells[i].paragraphs[0]
    run = p.add_run(h_text)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(t.rows[0].cells[i], "1F4E79")

scales = [
    ("15 users (current)", "t3.small, single instance", "£36-53"),
    ("30 users", "t3.medium, single instance", "£45-65"),
    ("50-100 users", "t3.large or load-balanced EB", "£70-120"),
    ("100+ users", "Migrate to ECS Fargate with auto-scaling", "£120-200+"),
]
for r_idx, (users, infra, cost) in enumerate(scales, 1):
    t.rows[r_idx].cells[0].paragraphs[0].add_run(users).font.size = Pt(9)
    t.rows[r_idx].cells[0].paragraphs[0].runs[0].font.bold = True
    t.rows[r_idx].cells[1].paragraphs[0].add_run(infra).font.size = Pt(9)
    run = t.rows[r_idx].cells[2].paragraphs[0].add_run(cost)
    run.font.size = Pt(9)
    run.font.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Key point: ")
run.font.bold = True
run = p.add_run(
    "You don't need to over-provision. Start with the smallest instance and "
    "upgrade only when needed. Elastic Beanstalk makes it easy to change instance "
    "types with zero downtime."
)
run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════
# 10. TIMELINE
# ═══════════════════════════════════════════
h = doc.add_heading("10. Timeline & Next Steps", level=1)
for run in h.runs:
    run.font.color.rgb = DARK_BLUE

t = doc.add_table(rows=6, cols=3)
t.style = 'Table Grid'

time_headers = ["Phase", "Effort", "Description"]
for i, h_text in enumerate(time_headers):
    p = t.rows[0].cells[i].paragraphs[0]
    run = p.add_run(h_text)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(t.rows[0].cells[i], "1F4E79")

timeline = [
    ("Phase 1", "1-2 days", "Fix critical code issues (auth, CORS, rate limiting, deps)"),
    ("Phase 2", "1 day", "Docker containerisation and local testing"),
    ("Phase 3", "1-2 days", "AWS infrastructure (EB, S3, CloudFront, DNS, SSL)"),
    ("Phase 4", "0.5 day", "WordPress integration (iframe embed)"),
    ("Phase 5", "0.5 day", "Monitoring, testing, go-live"),
]
for r_idx, (phase, effort, desc) in enumerate(timeline, 1):
    t.rows[r_idx].cells[0].paragraphs[0].add_run(phase).font.size = Pt(10)
    t.rows[r_idx].cells[0].paragraphs[0].runs[0].font.bold = True
    run = t.rows[r_idx].cells[1].paragraphs[0].add_run(effort)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = NHS_BLUE
    t.rows[r_idx].cells[2].paragraphs[0].add_run(desc).font.size = Pt(10)
    if r_idx % 2 == 0:
        for cell in t.rows[r_idx].cells:
            set_cell_shading(cell, "F5F5F5")

doc.add_paragraph()

total_p = doc.add_paragraph()
run = total_p.add_run("Total: 4-5 working days from start to live deployment")
run.font.bold = True
run.font.size = Pt(13)
run.font.color.rgb = GREEN

doc.add_paragraph()
doc.add_paragraph()

# Next steps
doc.add_heading("Immediate Next Steps", level=2)

next_steps = [
    "1. Team reviews this document and approves architecture",
    "2. Team reviews model comparison documents and decides on Nova Pro vs Sonnet",
    "3. Begin Phase 1 — fix critical code issues",
    "4. Register domain/subdomain for chatbot (chatbot.yoursite.com)",
    "5. Request ACM SSL certificate for the subdomain",
]
for step in next_steps:
    p = doc.add_paragraph()
    run = p.add_run(step)
    run.font.size = Pt(11)

# Save
output_path = "/Users/CajaLtd/Chatbot/Architecture_and_Costing.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
